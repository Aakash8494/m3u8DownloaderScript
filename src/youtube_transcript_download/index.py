import os
import re
from urllib.parse import urlparse, parse_qs
import yt_dlp
from youtube_transcript_api import YouTubeTranscriptApi
from docx import Document

def extract_video_id(url: str) -> str:
    """Extracts the YouTube Video ID from a standard, shorts, or embed URL."""
    url = url.strip()
    parsed = urlparse(url)
    if parsed.scheme in ("http", "https"):
        if parsed.netloc in ("youtu.be", "www.youtu.be"):
            return parsed.path.lstrip("/")
        qs = parse_qs(parsed.query)
        if "v" in qs:
            return qs["v"][0]
        m = re.search(r"/(shorts|embed)/([^/?&#]+)", parsed.path)
        if m:
            return m.group(2)
        return parsed.path.lstrip("/")
    return url

def get_single_video_info(url: str):
    """Uses yt-dlp to get the title and channel of a single video."""
    ydl_opts = {'quiet': True, 'extract_flat': True}
    with yt_dlp.YoutubeDL(ydl_opts) as ydl:
        info = ydl.extract_info(url, download=False)
        return {
            'id': info.get('id', extract_video_id(url)),
            'title': info.get('title', 'Unknown Title'),
            'channel': info.get('uploader', info.get('channel', 'Unknown Channel'))
        }

def get_playlist_videos(playlist_url: str):
    """
    Uses yt-dlp to extract all video IDs, titles, and channels from a playlist.
    Returns a tuple: (playlist_title, list_of_video_dicts)
    """
    parsed = urlparse(playlist_url)
    qs = parse_qs(parsed.query)
    if 'list' in qs:
        playlist_id = qs['list'][0]
        playlist_url = f"https://www.youtube.com/playlist?list={playlist_id}"

    ydl_opts = {
        'extract_flat': 'in_playlist',
        'quiet': True,
        'ignoreerrors': True
    }
    
    with yt_dlp.YoutubeDL(ydl_opts) as ydl:
        info = ydl.extract_info(playlist_url, download=False)
        if not info:
            return "Unknown Playlist", []
            
        playlist_title = info.get('title', 'Unknown_Playlist')
        videos = []
        
        if 'entries' in info:
            for entry in info['entries']:
                if entry and isinstance(entry, dict) and entry.get('id'):
                    videos.append({
                        'id': entry['id'],
                        'title': entry.get('title', 'Unknown Title'),
                        'channel': entry.get('uploader', entry.get('channel', 'Unknown Channel'))
                    })
        return playlist_title, videos

def fetch_transcript(video_id: str, languages=['hi', 'en']):
    """
    Fetches the transcript for a given video ID. 
    Prefers manual transcripts, falls back to auto-generated.
    """
    api = YouTubeTranscriptApi()
    try:
        transcript_list = api.list_transcripts(video_id)
        try:
            transcript_obj = transcript_list.find_transcript(languages)
        except:
            transcript_obj = transcript_list.find_generated_transcript(languages)
            
        raw_data = transcript_obj.fetch()
        text_lines = [entry['text'] for entry in raw_data]
        
    except Exception as e:
        try:
            transcript_obj = api.fetch(video_id, languages=languages)
            text_lines = [entry['text'] if isinstance(entry, dict) else entry.text for entry in transcript_obj]
        except Exception as fallback_err:
            print(f"❌ Failed to fetch transcript for {video_id}: {fallback_err}")
            return None

    return " ".join(text_lines)

def clean_name(name: str) -> str:
    """Removes invalid characters from folder and file names."""
    return re.sub(r'[\\/*?:"<>|]', "", name).strip()

def save_to_word_doc(content: str, channel: str, title: str):
    """
    Creates a folder for the channel and saves the transcript as a Word document.
    """
    safe_channel = clean_name(channel) if channel else "Unknown Channel"
    safe_title = clean_name(title) if title else "Unknown Video"
    
    # Create the channel directory in the current working folder
    channel_dir = os.path.join(os.getcwd(), safe_channel)
    os.makedirs(channel_dir, exist_ok=True)
    
    # Define the file path (Channel Name/Video Name.docx)
    file_path = os.path.join(channel_dir, f"{safe_title}.docx")
    
    # Create the Word document
    doc = Document()
    doc.add_heading(title, level=1)
    
    # Add metadata
    meta_para = doc.add_paragraph()
    meta_para.add_run(f"Channel: {channel}").italic = True
    
    # Add the transcript content
    doc.add_paragraph(content)
    
    # Save the document
    doc.save(file_path)
        
    return file_path

# ==========================================
# Example Usage
# ==========================================
if __name__ == "__main__":
    url = input("Enter a YouTube Video or Playlist URL: ").strip()
    
    is_playlist = "list=" in url
    
    if is_playlist:
        print("\n📂 Fetching playlist data...")
        playlist_title, videos = get_playlist_videos(url)
        print(f"Found {len(videos)} videos in '{playlist_title}'.\n")
        
        # NOTE: Removed the [:3] limit. This will now process the entire playlist.
        for vid in videos: 
            print(f"▶ Downloading transcript for: {vid['title']}")
            transcript = fetch_transcript(vid['id'])
            
            if transcript:
                saved_path = save_to_word_doc(transcript, vid['channel'], vid['title'])
                print(f"✅ Saved to: {saved_path}\n")
                
    else:
        print("\n🎥 Fetching single video data...")
        vid_info = get_single_video_info(url)
        print(f"Video Title: {vid_info['title']}")
        print(f"Channel: {vid_info['channel']}")
        
        transcript = fetch_transcript(vid_info['id'])
        if transcript:
            saved_path = save_to_word_doc(transcript, vid_info['channel'], vid_info['title'])
            print(f"\n✅ Transcript successfully downloaded!")
            print(f"📁 Saved to: {saved_path}")