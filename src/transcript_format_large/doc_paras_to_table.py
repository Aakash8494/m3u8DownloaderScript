import argparse
from pathlib import Path
from docx import Document

def is_title_paragraph(para):
    """
    Checks if a paragraph should be considered a Title/Heading.
    Returns True if it has a Heading style OR if the entire line is bold.
    """
    style_name = para.style.name.lower()
    if 'heading' in style_name or 'title' in style_name:
        return True
    
    # Fallback check: if all the text in the paragraph is bold, treat as title
    runs_with_text = [r for r in para.runs if r.text.strip()]
    if runs_with_text and all(r.bold for r in runs_with_text):
        return True
        
    return False

def extract_to_single_column(input_path: Path):
    print(f"Reading document: {input_path.name}...")
    
    try:
        doc = Document(input_path)
    except Exception as e:
        print(f"Error opening document: {e}")
        return

    # 1. Create the new document and table
    new_doc = Document()
    new_doc.add_heading(f'Extracted: {input_path.stem}', level=1)
    
    table = new_doc.add_table(rows=0, cols=1)
    table.style = 'Table Grid' 

    blocks_processed = 0
    current_cell = None

    # 2. Process every single paragraph one by one
    for para in doc.paragraphs:
        if not para.text.strip():
            continue # Skip blank lines

        # 3. Check if we need a new row or if we add to the current one
        if current_cell is None or is_title_paragraph(para):
            # Create a brand new row for this Title
            row = table.add_row()
            current_cell = row.cells[0]
            new_para = current_cell.paragraphs[0] # Word creates an empty para by default
            blocks_processed += 1
        else:
            # Not a title, so add a paragraph inside the SAME cell
            new_para = current_cell.add_paragraph()
        
        # 4. Attempt to preserve the paragraph's overall style
        try:
            new_para.style = para.style.name
        except KeyError:
            pass # Fallback to standard text if the new blank doc doesn't have that style

        # 5. Copy the text chunk-by-chunk (Run-by-Run) to preserve BOLD formatting
        for run in para.runs:
            if run.text:
                new_run = new_para.add_run(run.text)
                new_run.bold = run.bold
                new_run.italic = run.italic
                # You can also add new_run.underline = run.underline here if needed

    print(f"Processed {blocks_processed} grouped blocks into the single-column table.")

    # 6. Generate dynamic output path and save
    output_path = input_path.parent / f"{input_path.stem}_SingleCol{input_path.suffix}"
    
    new_doc.save(output_path)
    print(f"\n✅ Success! Saved formatted table to:\n📁 {output_path}")

if __name__ == "__main__":
    # Setup Command Line Interface
    parser = argparse.ArgumentParser(description="Convert Word doc into a 1-column table while grouping titles and text.")
    parser.add_argument("filepath", type=str, help="Full path to the .docx file")
    
    args = parser.parse_args()
    
    # Clean the path
    clean_path_str = args.filepath.strip('"').strip("'")
    input_path = Path(clean_path_str).resolve()
    
    # Validate the file
    if not input_path.exists():
        print(f"❌ Error: Could not find file at '{input_path}'")
    elif input_path.suffix.lower() != '.docx':
        print(f"❌ Error: The file must be a standard Word Document (.docx). You provided: '{input_path.suffix}'")
    else:
        extract_to_single_column(input_path)