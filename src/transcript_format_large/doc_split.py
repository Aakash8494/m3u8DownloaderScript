from docx import Document
import math
import os
import sys

def split_docx_into_parts(file_path, n, output_dir="output_parts"):
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"File not found: {file_path}")

    doc = Document(file_path)

    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    full_text = "\n".join(paragraphs)

    words = full_text.split()
    total_words = len(words)

    if n <= 0:
        raise ValueError("n must be greater than 0")

    part_size = math.ceil(total_words / n)

    os.makedirs(output_dir, exist_ok=True)

    base_name = os.path.splitext(os.path.basename(file_path))[0]

    for idx, i in enumerate(range(0, total_words, part_size)):
        part_words = words[i:i + part_size]
        part_text = " ".join(part_words)

        new_doc = Document()
        for line in part_text.split("\n"):
            new_doc.add_paragraph(line)

        output_path = os.path.join(output_dir, f"{base_name}_part_{idx+1}.docx")
        new_doc.save(output_path)
        print(f"Saved: {output_path}")


# ✅ CLI ENTRY POINT
if __name__ == "__main__":
    if len(sys.argv) < 3:
        print("Usage: python doc_split.py <file_path> <n_parts>")
        sys.exit(1)

    file_path = sys.argv[1]
    n = int(sys.argv[2])

    split_docx_into_parts(file_path, n)