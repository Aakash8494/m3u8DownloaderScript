import os
import re
import time
import traceback
from dotenv import load_dotenv
from docx import Document

# --- OPENAI SDK ---
from openai import OpenAI
# ------------------

# 1. Force Python to use the exact folder where this script is located
script_dir = os.path.dirname(os.path.abspath(__file__))
os.chdir(script_dir)

# Load environment variables
load_dotenv()

# Safely retrieve API Key
api_key = os.getenv("OPENAI_API_KEY")
if not api_key:
    print("❌ Error: OPENAI_API_KEY not found. Please ensure your .env file is in this folder.")
    input("\nPress Enter to exit...")
    exit(1)

# Initialize OpenAI client
client = OpenAI(api_key=api_key)

# ==========================================
# PROMPTS
# ==========================================
PROMPT = """
You are an expert transcriber, strict text formatter, and translator. I am providing you with a raw, unpunctuated text transcript from a video.

CRITICAL LANGUAGE AND SCRIPT RULE:
- If the provided text is in Hindi (written in Devanagari script, e.g., नमस्ते), you MUST transliterate it completely into Hinglish (Hindi words written using the English alphabet, e.g., namaste). 
- The final output MUST NOT contain ANY Devanagari/Hindi script characters. It must be written 100% using the English alphabet.
- If the text is in English, leave it exactly in English.

You must process the exact spoken words verbatim. Do NOT add, delete, summarize, or skip a single spoken word from the provided text. (Note: Transliterating the script from Hindi to English letters does not violate this rule; you must preserve the exact spoken words, just change the alphabet).

Since the raw text lacks visible punctuation, you must deduce and add periods (full stops) where sentences naturally end. Always capitalize the first letter of the word immediately following a period.

Break the continuous text down into readable paragraphs of roughly 150 to 200 words. Additionally, if the topic clearly transitions, start a new paragraph.

Create a short, appropriate title/heading for every single paragraph based on what that section is about. Place the title right above its matching paragraph.

Find the most important keywords, phrases, or main ideas inside each formatted paragraph and make them bold to make it easier to read.

Return only the final formatted text. Do not add any extra introductory or concluding sentences before or after the text.
"""

TITLE_PROMPT = """
If the following text contains any Hindi/Devanagari script, transliterate it completely into Hinglish (Hindi words written using the English alphabet). 
If it is already entirely in English, return it exactly as is. 
Only return the transformed title text. Do not add quotes, introductions, or extra punctuation.
"""

# ==========================================
# HELPER FUNCTIONS
# ==========================================
def read_docx(file_path: str) -> str:
    doc = Document(file_path)
    full_text = []
    for para in doc.paragraphs:
        if para.text.strip():
            full_text.append(para.text.strip())
    return " ".join(full_text)

def safe_add_heading(doc, text, level):
    try:
        doc.add_heading(text, level=level)
    except KeyError:
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = True

def create_formatted_doc(formatted_text: str, original_filename: str, output_path: str):
    doc = Document()

    safe_add_heading(doc, f"Formatted Transcript: {original_filename}", level=1)
    doc.add_paragraph()

    lines = formatted_text.split('\n')

    for line in lines:
        text = line.strip()
        if not text:
            continue

        if text.startswith("### "):
            safe_add_heading(doc, text[4:].strip(), level=3)
        elif text.startswith("## "):
            safe_add_heading(doc, text[3:].strip(), level=2)
        elif text.startswith("# "):
            safe_add_heading(doc, text[2:].strip(), level=1)
        else:
            para = doc.add_paragraph()
            parts = text.split("**")
            for idx, part in enumerate(parts):
                if part:
                    run = para.add_run(part)
                    if idx % 2 != 0:
                        run.bold = True

    doc.save(output_path)

def generate_with_retry(prompt_text, text_content, max_retries=5):
    retry_delay = 5

    for attempt in range(max_retries):
        try:
            response = client.chat.completions.create(
                model="gpt-4.1-mini",
                messages=[
                    {"role": "system", "content": prompt_text},
                    {"role": "user", "content": text_content}
                ],
                temperature=0.2,
                max_tokens=4000
            )

            return response.choices[0].message.content

        except Exception as e:
            error_msg = str(e)

            if "429" in error_msg or "503" in error_msg:
                if attempt < max_retries - 1:
                    print(f"⚠️ Server busy (Attempt {attempt + 1}/{max_retries}). Retrying in {retry_delay} seconds...")
                    time.sleep(retry_delay)
                    retry_delay *= 2
                    continue

            raise e

    return None

# ==========================================
# MAIN EXECUTION
# ==========================================
def main():
    print("================================================")
    print("  Local Document Transcriber & Formatter (ChatGPT)")
    print("================================================\n")

    file_path = input("📄 Enter the full path to your .docx file:\n> ").strip()
    file_path = file_path.strip('"').strip("'")

    if not os.path.exists(file_path):
        print("\n❌ Error: File not found.")
        return

    if not file_path.lower().endswith('.docx'):
        print("\n❌ Error: Please provide a .docx file.")
        return

    dir_name = os.path.dirname(file_path)
    base_name = os.path.basename(file_path)
    name_without_ext = os.path.splitext(base_name)[0]

    print(f"\n📖 Reading document: {base_name}...")
    try:
        raw_text = read_docx(file_path)
        if not raw_text.strip():
            print("❌ The document appears to be empty.")
            return
        print(f"✅ Read {len(raw_text.split())} words.")
    except Exception as e:
        print(f"❌ Failed to read document: {e}")
        return

    try:
        print("🤖 Cleaning title...")
        clean_title = generate_with_retry(TITLE_PROMPT, f"Title to clean: {name_without_ext}", max_retries=3)
        if not clean_title:
            clean_title = name_without_ext
    except Exception:
        clean_title = name_without_ext

    clean_title = re.sub(r'[\\/*?:"<>|]', "", clean_title.strip())
    output_filename = f"{clean_title}_Formatted.docx"
    output_path = os.path.join(dir_name, output_filename)

    print("\n🤖 Sending text to ChatGPT for formatting...")
    try:
        formatted_text = generate_with_retry(PROMPT, f"Transcript: {raw_text}", max_retries=5)

        if not formatted_text:
            print("❌ Empty response from ChatGPT.")
            return

    except Exception as e:
        print(f"\n❌ API error: {e}")
        return

    print("\n📝 Generating new Word document...")
    try:
        create_formatted_doc(formatted_text, clean_title, output_path)
        print("✅ SUCCESS! Document saved:")
        print(f"📁 {output_path}")
    except Exception as e:
        print(f"❌ Failed to save document: {e}")

if __name__ == "__main__":
    try:
        main()
    except KeyboardInterrupt:
        print("\n\n🛑 Cancelled by user.")
    except Exception:
        print("\n" + "="*50)
        print("❌ FATAL ERROR:")
        print("="*50)
        traceback.print_exc()
        print("="*50)

    input("\nPress Enter to exit...")