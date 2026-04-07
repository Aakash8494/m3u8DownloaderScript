import os
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# =========================================================
# --- FORMATTING HELPER FUNCTIONS ---
# =========================================================
def add_page_number(run):
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE"
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    run._r.extend([fldChar1, instrText, fldChar2, fldChar3])

def apply_custom_formatting(doc):
    """Applies your custom margins, fonts, and strictly disables page breaks."""
    for section in doc.sections:
        section.top_margin = Inches(0.0)
        section.bottom_margin = Inches(0.0)
        section.left_margin = Inches(0.17)
        section.right_margin = Inches(4.0)
        
        footer = section.footer
        footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_page_number(footer_para.add_run())

    doc.styles['Normal'].font.size = Pt(8)
    
    # Apply heading sizes and explicitly disable "Page Break Before" to save paper
    fonts = {'Heading 1': 14, 'Heading 2': 12, 'Heading 3': 10}
    for style_name, size in fonts.items():
        try:
            style = doc.styles[style_name]
            style.font.size = Pt(size)
            style.paragraph_format.page_break_before = False
        except KeyError:
            pass

def copy_doc_content(source_path, target_doc):
    """Copies all paragraphs and formatting from an individual doc into the master doc."""
    source_doc = Document(source_path)
    for para in source_doc.paragraphs:
        new_para = target_doc.add_paragraph(style=para.style.name)
        new_para.alignment = para.alignment
        for run in para.runs:
            new_run = new_para.add_run(run.text)
            new_run.bold = run.bold
            new_run.italic = run.italic
            new_run.underline = run.underline

# =========================================================
# --- CORE MERGE LOGIC ---
# =========================================================
def merge_docs_in_folder(folder_path, output_filename="Master_Compiled_Docs.docx"):
    folder_path = os.path.abspath(folder_path)

    if not os.path.isdir(folder_path):
        print(f"❌ Error: The directory '{folder_path}' does not exist.")
        return

    print(f"📂 Scanning folder: {folder_path}")
    
    # Find all docx files, ignoring the master file itself and temp word files (~$)
    docx_files = []
    for f in os.listdir(folder_path):
        if f.lower().endswith('.docx') and not f.startswith('~$') and f != output_filename:
            docx_files.append(f)
            
    if not docx_files:
        print("⚠️ No valid .docx files found to merge.")
        return

    # Sort them alphabetically so they appear in a predictable order
    docx_files.sort()
    print(f"📄 Found {len(docx_files)} files to merge.")

    # Initialize the master document with your custom layout
    master_doc = Document()
    apply_custom_formatting(master_doc)

    # Loop through and append each file
    for i, filename in enumerate(docx_files):
        file_path = os.path.join(folder_path, filename)
        print(f"  [{i+1}/{len(docx_files)}] Appending: {filename}...")
        
        try:
            copy_doc_content(file_path, master_doc)
        except Exception as e:
            print(f"  ❌ Error reading {filename}: {e}")
            continue

        # Add a single blank line between different documents
        if i < len(docx_files) - 1:
            master_doc.add_paragraph()

    output_path = os.path.join(folder_path, output_filename)
    
    try:
        master_doc.save(output_path)
        print(f"\n✅ Success! Saved master document to:\n📁 {output_path}")
    except PermissionError:
        print(f"\n❌ Permission Error: Could not save to {output_path}.")
        print("Is the 'Master_Compiled_Docs.docx' file currently open in Microsoft Word? If so, close it and run this again.")

if __name__ == "__main__":
    print("\n" + "=" * 50)
    print("📑 Word Document Auto-Merger 📑")
    print("=" * 50)
    
    folder_input = input("Enter the path to the folder containing your .docx files\n(or just press Enter to scan the current folder): ").strip()
    folder_path = folder_input if folder_input else "."
    
    print("\nStarting process...\n")
    merge_docs_in_folder(folder_path)
    
    input("\nAll tasks finished. Press Enter to exit...")