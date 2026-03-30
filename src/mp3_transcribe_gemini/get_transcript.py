import os
import re
import time
import requests
import traceback
import tkinter as tk
from urllib.parse import urlparse, parse_qs
from youtube_transcript_api import YouTubeTranscriptApi, TranscriptsDisabled, NoTranscriptFound

import google.generativeai as genai

# --- NEW DOCX IMPORTS ---
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
# ------------------------

from dotenv import load_dotenv

# 1. Force Python to use the exact folder where this script is located
script_dir = os.path.dirname(os.path.abspath(__file__))
os.chdir(script_dir)

# Load environment variables from the shared .env file
load_dotenv()

# Safely retrieve API Key
api_key = os.getenv("GEMINI_API_KEY")
if not api_key:
    print("❌ Error: GEMINI_API_KEY not found. Please ensure your .env file is in this folder.")
    input("\nPress Enter to exit...")
    exit(1)

genai.configure(api_key=api_key)
model = genai.GenerativeModel('gemini-2.5-flash')

# Define the Prompt for the main transcript
PROMPT = """
You are an expert transcriber, strict text formatter, and translator. I am providing you with a raw, unpunctuated text transcript from a video.

CRITICAL LANGUAGE AND SCRIPT RULE:
- If the provided text is in Hindi (written in Devanagari script, e.g., नमस्ते), you MUST transliterate it completely into Hinglish (Hindi words written using the English alphabet, e.g., namaste). 
- The final output MUST NOT contain ANY Devanagari/Hindi script characters. It must be written 100% using the English alphabet.
- If the text is in English, leave it exactly in English.

You must process the exact spoken words verbatim. Do NOT add, delete, summarize, or skip a single spoken word from the provided text. (Note: Transliterating the script from Hindi to English letters does not violate this rule; you must preserve the exact spoken words, just change the alphabet).

Since the raw text lacks visible punctuation, you must deduce and add periods (full stops) where sentences naturally end. Always capitalize the first letter of the word immediately following a period.

Break the continuous text down into readable paragraphs of roughly 150 to 200 words. Additionally, if the topic clearly transitions, start a new paragraph.

Create a short, appropriate title/heading for every single paragraph based on what that section is about. Place the title right above its matching paragraph.

Find the most important keywords, phrases, or main ideas inside each formatted paragraph and make them bold to make it easier to read.

Return only the final formatted text. Do not add any extra introductory or concluding sentences before or after the text.
"""

# Define a mini-prompt just for cleaning up the Title
TITLE_PROMPT = """
If the following text contains any Hindi/Devanagari script, transliterate it completely into Hinglish (Hindi words written using the English alphabet). 
If it is already entirely in English, return it exactly as is. 
Only return the transformed title text. Do not add quotes, introductions, or extra punctuation.
"""

# --- NEW HELPER FUNCTION FOR PAGE NUMBERS ---
def add_page_number(run):
    """Injects Word XML field codes to auto-generate page numbers."""
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')

    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE"

    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')

    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')

    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(fldChar3)
# --------------------------------------------

def get_clipboard_text():
    """Silently grabs whatever text is currently copied to the clipboard."""
    root = tk.Tk()
    root.withdraw()
    try:
        text = root.clipboard_get()
    except tk.TclError:
        text = ""
    root.destroy()
    return text

def extract_video_id(value: str) -> str:
    value = value.strip()
    parsed = urlparse(value)
    if parsed.scheme in ("http", "https"):
        if parsed.netloc in ("youtu.be", "www.youtu.be"):
            return parsed.path.lstrip("/")
        qs = parse_qs(parsed.query)
        if "v" in qs:
            return qs["v"][0]
        m = re.search(r"/(shorts|embed)/([^/?&#]+)", parsed.path)
        if m:
            return m.group(2)
        return parsed.path.lstrip("/")
    return value

def clean_filename(name: str) -> str:
    return re.sub(r'[\\/*?:"<>|]', "", name).strip()

def get_title_and_channel(video_page_url: str):
    oembed_url = "https://www.youtube.com/oembed"
    params = {"url": video_page_url, "format": "json"}
    try:
        resp = requests.get(oembed_url, params=params, timeout=10)
        resp.raise_for_status()
        data = resp.json()
        return data.get("title", "Untitled"), data.get("author_name", "UnknownChannel")
    except:
        return "Unknown_Video", "Unknown_Channel"

def safe_add_heading(doc, text, level):
    """Safely adds a heading, falling back to bold text if the style is missing."""
    try:
        doc.add_heading(text, level=level)
    except KeyError:
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = True

def main():
    video_input = get_clipboard_text().strip()

    if "youtube.com" not in video_input and "youtu.be" not in video_input:
        print("❌ Error: No YouTube URL found in your clipboard.")
        print(f"What we found instead: '{video_input[:50]}...'")
        input("\nPlease copy a YouTube link and try again. Press Enter to exit...")
        return

    print(f"📋 Read from clipboard: {video_input}\n")

    try:
        video_id = extract_video_id(video_input)
        page_url = f"https://www.youtube.com/watch?v={video_id}"

        print("Fetching video metadata...")
        raw_title, channel = get_title_and_channel(page_url)
        print(f"Raw Title Found: {raw_title} ({channel})")

        print("Checking title for Hindi script...")
        try:
            title_response = model.generate_content([TITLE_PROMPT, raw_title])
            title = title_response.text.strip()
            if title != raw_title:
                print(f"Transliterated Title: {title}")
        except Exception as e:
            print(f"Title translation failed (using raw): {e}")
            title = raw_title

        clean_channel = clean_filename(channel)
        clean_title = clean_filename(title)

        base_dir = "YouTube_Transcripts"
        channel_dir = os.path.join(base_dir, clean_channel)
        os.makedirs(channel_dir, exist_ok=True)

        txt_file_path = os.path.join(channel_dir, f"{clean_title}.txt")
        doc_file_path = os.path.join(channel_dir, f"{clean_title}.docx")

        print("Downloading raw transcript...")
        api = YouTubeTranscriptApi()
        
        try:
            transcript_list = api.list_transcripts(video_id)
            try:
                transcript_obj = transcript_list.find_transcript(['hi', 'en'])
            except:
                transcript_obj = transcript_list.find_generated_transcript(['hi', 'en'])
                
            raw_data = transcript_obj.fetch()
            text_lines = [entry['text'] for entry in raw_data]
            
        except Exception as transcript_err:
             print(f"Transcript fetch logic failed, trying fallback: {transcript_err}")
             transcript_obj = api.fetch(video_id, languages=['hi', 'en'])
             text_lines = [entry['text'] if isinstance(entry, dict) else entry.text for entry in transcript_obj]

        video_link = f"https://www.youtube.com/watch?v={video_id}"
        header = f"Title:   {title}\nChannel: {channel}\nLink:    {video_link}\n{'-' * 50}\n\n"
        raw_full_text = " ".join(text_lines)
        
        with open(txt_file_path, "w", encoding="utf-8") as f:
            f.write(header + raw_full_text)
        print(f"✅ Saved raw backup: {txt_file_path}")

        print("🤖 Sending raw text to Gemini for formatting and translation...")
        response = model.generate_content([PROMPT, raw_full_text])
        formatted_text = response.text

        print("📝 Creating beautifully formatted Word Document...")
        doc = Document()
        
        # =========================================================
        # --- NEW CUSTOM DOCUMENT FORMATTING ---
        # =========================================================
        
        # 1. Set Custom Margins (Top: 0, Bottom: 0, Left: 0.17", Right: 4")
        for section in doc.sections:
            section.top_margin = Inches(0)
            section.bottom_margin = Inches(0)
            section.left_margin = Inches(0.17)
            section.right_margin = Inches(4.0)
            
            # 2. Add Page Numbers to Footer Center
            footer = section.footer
            footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
            footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = footer_para.add_run()
            add_page_number(run)

        # 3. Shrink overall fonts 3 times
        # Default font is 11pt, 3 shrinks = 8pt
        style_normal = doc.styles['Normal']
        style_normal.font.size = Pt(8)
        
        # Also proportionally shrink the Headings so they don't look giant 
        # compared to the 8pt body text.
        try:
            doc.styles['Heading 1'].font.size = Pt(14)
            doc.styles['Heading 2'].font.size = Pt(12)
            doc.styles['Heading 3'].font.size = Pt(10)
        except KeyError:
            pass
            
        # =========================================================

        # Add title
        safe_add_heading(doc, f"Source: {title}", level=1)
        
        lines = formatted_text.split('\n')
        
        for line in lines:
            text = line.strip()
            if not text:
                continue 
                
            if text.startswith("### "):
                safe_add_heading(doc, text[4:].strip(), level=3)
            elif text.startswith("## "):
                safe_add_heading(doc, text[3:].strip(), level=2)
            elif text.startswith("# "):
                safe_add_heading(doc, text[2:].strip(), level=1)
            else:
                para = doc.add_paragraph()
                parts = text.split("**")
                
                for idx, part in enumerate(parts):
                    if part: 
                        run = para.add_run(part)
                        if idx % 2 != 0:
                            run.bold = True

        doc.save(doc_file_path)
        print(f"✅ Saved formatted Word doc: {doc_file_path}")

    except NoTranscriptFound:
        print("❌ Error: No transcript found in Hindi or English.")
    except TranscriptsDisabled:
        print("❌ Error: Transcripts are disabled for this video.")
    except Exception as e:
        print(f"❌ Error: {e}")

    input("\nPress Enter to close this window...")

if __name__ == "__main__":
    try:
        main()
    except Exception as e:
        print("\n" + "="*50)
        print("❌ A FATAL ERROR OCCURRED:")
        print("="*50)
        traceback.print_exc()
        print("="*50)
        input("\nPress Enter to close this window...")