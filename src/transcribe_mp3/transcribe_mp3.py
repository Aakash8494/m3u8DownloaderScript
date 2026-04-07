import os
import re
import time
import json
import logging
import concurrent.futures
from functools import partial

# --- NEW GEMINI SDK IMPORTS ---
from google import genai
# ------------------------------

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from dotenv import load_dotenv

# =========================================================
# --- CONFIGURATION & LOGGING ---
# =========================================================
CONFIG_FILE = "transcriber_config.json"

DEFAULT_CONFIG = {
    "MAX_WORKERS": 1, # FORCED TO 1 TO PREVENT RATE LIMITS
    "MAX_RETRIES": 5, # INCREASED TO ALLOW FOR LONG COOLDOWNS
    "MARGINS": {"top": 0.0, "bottom": 0.0, "left": 0.17, "right": 4.0},
    "FONT_SIZES": {"Normal": 8, "Heading 1": 14, "Heading 2": 12, "Heading 3": 10},
    "PROMPT": (
        "You are an expert audio transcriber, strict text formatter, and translator. I have provided an MP3 audio file containing speech in either Hindi or English.\n\n"
        "If the spoken audio is in Hindi, transcribe and transliterate it directly into Hinglish (Hindi words written using the English alphabet). If the spoken audio is in English, transcribe it exactly in English.\n"
        "You must transcribe the exact spoken words. Do NOT add, delete, summarize, or skip a single spoken word from the audio. It must be a 100% word-for-word verbatim transcription of the speaker.\n\n"
        "Since the audio lacks visible punctuation, you must deduce and add periods (full stops) where sentences naturally end based on the speaker's pauses and intonation. Always capitalize the first letter of the word immediately following a period.\n\n"
        "Break the continuous transcription down into readable paragraphs of roughly 150 to 200 words. Additionally, if the speaker takes a significant pause or clearly transitions to a new topic, start a new paragraph.\n\n"
        "Create a short, appropriate title/heading for every single paragraph based on what that section of the audio is about. Place the title right above its matching paragraph.\n\n"
        "Find the most important keywords, phrases, or main ideas inside each transcribed paragraph and make them bold to make it easier to read.\n\n"
        "Return only the final formatted transcription. Do not add any extra introductory or concluding sentences before or after the text."
    )
}

def load_config():
    if not os.path.exists(CONFIG_FILE):
        with open(CONFIG_FILE, 'w', encoding='utf-8') as f:
            json.dump(DEFAULT_CONFIG, f, indent=4)
        return DEFAULT_CONFIG
    with open(CONFIG_FILE, 'r', encoding='utf-8') as f:
        return json.load(f)

CONFIG = load_config()
CONFIG["MAX_WORKERS"] = 1
CONFIG["MAX_RETRIES"] = 5

logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s [%(levelname)s] %(message)s",
    handlers=[
        logging.FileHandler("transcription.log", encoding='utf-8'),
        logging.StreamHandler()
    ]
)

load_dotenv()
api_key = os.getenv("GEMINI_API_KEY")
if not api_key:
    logging.error("GEMINI_API_KEY not found. Please ensure your .env file is set up correctly.")
    exit(1)

client = genai.Client(api_key=api_key)

# =========================================================
# --- FILENAME SANITIZATION HELPER ---
# =========================================================
def sanitize_filename(filename):
    """STRICT CLEANUP: Keeps only English letters, numbers, and spaces."""
    name, ext = os.path.splitext(filename)
    
    # Nuke everything that is NOT an uppercase letter, lowercase letter, number, or space
    name = re.sub(r'[^a-zA-Z0-9\s]', '', name)
    
    # Remove any awkward double spaces left behind by the deleted characters
    name = re.sub(r'\s+', ' ', name).strip()
    
    # Fallback if the entire filename was emojis or Hindi script and got completely wiped
    if not name:
        name = f"audio_track_{int(time.time())}"
        
    return name + ext

def sanitize_directory_filenames(folder_path):
    """Sweeps the directory and permanently cleans up MP3 and DOCX filenames."""
    for root, dirs, files in os.walk(folder_path):
        for filename in files:
            if filename.lower().endswith(('.mp3', '.docx')):
                clean_name = sanitize_filename(filename)
                if clean_name != filename:
                    old_path = os.path.join(root, filename)
                    new_path = os.path.join(root, clean_name)
                    
                    # Ensure we don't overwrite another file
                    if os.path.exists(new_path):
                        base, ext = os.path.splitext(clean_name)
                        new_path = os.path.join(root, f"{base}_{int(time.time())}{ext}")
                        
                    try:
                        os.rename(old_path, new_path)
                        logging.info(f"🧹 Sanitized filename: '{filename}' -> '{os.path.basename(new_path)}'")
                    except Exception as e:
                        logging.error(f"Failed to rename '{filename}': {e}")

# =========================================================
# --- FORMATTING HELPER FUNCTIONS ---
# =========================================================
def add_page_number(run):
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE"
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    run._r.extend([fldChar1, instrText, fldChar2, fldChar3])

def apply_custom_formatting(doc):
    margins = CONFIG["MARGINS"]
    for section in doc.sections:
        section.top_margin = Inches(margins["top"])
        section.bottom_margin = Inches(margins["bottom"])
        section.left_margin = Inches(margins["left"])
        section.right_margin = Inches(margins["right"])
        
        footer = section.footer
        footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        add_page_number(footer_para.add_run())

    fonts = CONFIG["FONT_SIZES"]
    doc.styles['Normal'].font.size = Pt(fonts["Normal"])
    
    for style_name in ['Heading 1', 'Heading 2', 'Heading 3']:
        try:
            style = doc.styles[style_name]
            style.font.size = Pt(fonts[style_name])
            style.paragraph_format.page_break_before = False
        except KeyError:
            pass

def append_transcription_to_doc(doc, raw_text):
    lines = raw_text.split('\n')
    for line in lines:
        text = line.strip()
        if not text:
            continue 
            
        if text.startswith("### "):
            doc.add_heading(text[4:].strip(), level=3)
        elif text.startswith("## "):
            doc.add_heading(text[3:].strip(), level=2)
        elif text.startswith("# "):
            doc.add_heading(text[2:].strip(), level=1)
        else:
            para = doc.add_paragraph()
            parts = text.split("**")
            for idx, part in enumerate(parts):
                if part: 
                    run = para.add_run(part)
                    if idx % 2 != 0:
                        run.bold = True

def copy_doc_content(source_path, target_doc):
    source_doc = Document(source_path)
    for para in source_doc.paragraphs:
        new_para = target_doc.add_paragraph(style=para.style.name)
        new_para.alignment = para.alignment
        for run in para.runs:
            new_run = new_para.add_run(run.text)
            new_run.bold = run.bold
            new_run.italic = run.italic
            new_run.underline = run.underline

# =========================================================
# --- CORE PROCESSING LOGIC ---
# =========================================================
def process_audio_task(file_path, current_folder):
    filename = os.path.basename(file_path)
    individual_filename = f"{os.path.splitext(filename)[0]}_transcribed.docx"
    individual_path = os.path.join(current_folder, individual_filename)
    expected_title = f"Source: {filename}"
    
    # --- 1. Check if Doc Already Exists ---
    if os.path.exists(individual_path):
        try:
            doc = Document(individual_path)
            if len(doc.paragraphs) > 0 and expected_title not in doc.paragraphs[0].text:
                logging.info(f"[{filename}] Existing doc found, but missing title. Injecting title...")
                doc.paragraphs[0].insert_paragraph_before(expected_title, style='Heading 1')
                doc.save(individual_path)
                return filename, individual_path, True, None
            else:
                logging.info(f"[{filename}] Existing doc found and valid. Skipping API.")
                return filename, individual_path, False, None 
        except Exception as e:
            return filename, None, False, f"Failed reading existing doc: {e}"

    # --- 2. Proceed with API Call if Doc is Missing ---
    logging.info(f"[{filename}] No doc found. Starting API processing...")

    raw_text = None
    error_msg = None
    
    for attempt in range(CONFIG["MAX_RETRIES"]):
        audio_file = None
        try:
            # Upload using the now completely stripped and safe filename
            audio_file = client.files.upload(file=file_path)
            
            while audio_file.state.name == "PROCESSING":
                time.sleep(2)
                audio_file = client.files.get(name=audio_file.name)
            
            logging.info(f"[{filename}] Upload complete. Generating text...")
            
            response = client.models.generate_content(
                model='gemini-2.5-flash',
                contents=[CONFIG["PROMPT"], audio_file]
            )
            raw_text = response.text
            break  
            
        except Exception as e:
            error_str = str(e)
            logging.warning(f"[{filename}] Attempt {attempt + 1} failed: {error_str}")
            
            if attempt == CONFIG["MAX_RETRIES"] - 1:
                error_msg = error_str
            else:
                # 🚨 SMART COOLDOWN FOR RATE LIMITS 🚨
                if "429" in error_str or "RESOURCE_EXHAUSTED" in error_str or "Quota" in error_str:
                    logging.info(f"⏳ Rate limit detected for [{filename}]. Sleeping for 65 seconds...")
                    time.sleep(65)
                else:
                    time.sleep(10)
                
        finally:
            if audio_file:
                try:
                    client.files.delete(name=audio_file.name)
                except Exception as cleanup_error:
                    logging.error(f"[{filename}] Failed to clean cloud file: {cleanup_error}")

    if error_msg:
        return filename, None, False, error_msg

    # --- 3. Save New Individual Document ---
    try:
        individual_doc = Document()
        apply_custom_formatting(individual_doc)
        individual_doc.add_heading(expected_title, level=1)
        append_transcription_to_doc(individual_doc, raw_text)
        individual_doc.save(individual_path)
        logging.info(f"[{filename}] Saved new individual document.")
        
        return filename, individual_path, True, None
        
    except Exception as e:
        return filename, None, False, f"Word formatting error: {str(e)}"

def process_and_merge_mp3s(folder_path, output_filename):
    folder_path = os.path.abspath(folder_path)

    if not os.path.isdir(folder_path):
        logging.error(f"The directory {folder_path} does not exist.")
        return
        
    # --- SANITIZE ALL FILES BEFORE DOING ANYTHING ELSE ---
    logging.info("Scanning and aggressively sanitizing filenames...")
    sanitize_directory_filenames(folder_path)

    files_by_folder = {}
    for root, dirs, files in os.walk(folder_path):
        mp3_files = [f for f in files if f.lower().endswith('.mp3')]
        if mp3_files:
            mp3_files.sort()
            files_by_folder[root] = [os.path.join(root, f) for f in mp3_files]
    
    if not files_by_folder:
        logging.warning(f"No MP3 files found in {folder_path} or its sub-directories.")
        return

    for current_folder, mp3_files_paths in files_by_folder.items():
        rel_path = os.path.relpath(current_folder, folder_path)
        display_dir = "main folder" if rel_path == "." else rel_path
        
        logging.info(f"=== Processing Directory: [{display_dir}] ({len(mp3_files_paths)} files) ===")

        task = partial(process_audio_task, current_folder=current_folder)
        with concurrent.futures.ThreadPoolExecutor(max_workers=CONFIG["MAX_WORKERS"]) as executor:
            results = list(executor.map(task, mp3_files_paths))

        output_path = os.path.join(current_folder, output_filename)
        master_needs_update = not os.path.exists(output_path) or any(updated for _, _, updated, _ in results)

        if master_needs_update:
            logging.info(f"Building/Updating continuous Master Document for [{display_dir}]...")
            master_doc = Document()
            apply_custom_formatting(master_doc)

            for i, (filename, ind_path, _, error_msg) in enumerate(results):
                if error_msg or not ind_path:
                    logging.error(f"Skipping {filename} in master doc due to failure: {error_msg}")
                    continue
                
                copy_doc_content(ind_path, master_doc)

                if i < len(results) - 1:
                    master_doc.add_paragraph()

            master_doc.save(output_path)
            logging.info(f"=== Success! Master doc for [{display_dir}] saved ===")
        else:
            logging.info(f"=== Master doc for [{display_dir}] already exists and is up-to-date. ===")

if __name__ == "__main__":
    print("\n" + "=" * 50)
    print("🎙️ Local Audio Transcriber & Formatter 🎙️")
    print("=" * 50)
    
    folder_input = input("Enter the path to the folder containing your MP3 files\n(or just press Enter to scan the current folder): ").strip()
    
    folder_path = folder_input if folder_input else "."
    output_filename = "Master_Transcriptions.docx"
    
    print("\nStarting process...\n")
    process_and_merge_mp3s(folder_path, output_filename)
    
    input("\nAll tasks finished. Press Enter to exit...")