import os
import re
import time
import requests
import traceback
import tkinter as tk
from urllib.parse import urlparse, parse_qs
from youtube_transcript_api import YouTubeTranscriptApi, TranscriptsDisabled, NoTranscriptFound

import google.generativeai as genai
import yt_dlp  # --- NEW IMPORT FOR PLAYLIST HANDLING ---

# --- NEW DOCX IMPORTS ---
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
# ------------------------

from dotenv import load_dotenv

# 1. Force Python to use the exact folder where this script is located
script_dir = os.path.dirname(os.path.abspath(__file__))
os.chdir(script_dir)

# Load environment variables from the shared .env file
load_dotenv()

# Safely retrieve API Key
api_key = os.getenv("GEMINI_API_KEY")
if not api_key:
    print("❌ Error: GEMINI_API_KEY not found. Please ensure your .env file is in this folder.")
    input("\nPress Enter to exit...")
    exit(1)

genai.configure(api_key=api_key)
model = genai.GenerativeModel('gemini-2.5-flash')

# Define the Prompt for the main transcript
PROMPT = """
You are an expert transcriber, strict text formatter, and translator. I am providing you with a raw, unpunctuated text transcript from a video.

CRITICAL LANGUAGE AND SCRIPT RULE:
- If the provided text is in Hindi (written in Devanagari script, e.g., नमस्ते), you MUST transliterate it completely into Hinglish (Hindi words written using the English alphabet, e.g., namaste). 
- The final output MUST NOT contain ANY Devanagari/Hindi script characters. It must be written 100% using the English alphabet.
- If the text is in English, leave it exactly in English.

You must process the exact spoken words verbatim. Do NOT add, delete, summarize, or skip a single spoken word from the provided text. (Note: Transliterating the script from Hindi to English letters does not violate this rule; you must preserve the exact spoken words, just change the alphabet).

Since the raw text lacks visible punctuation, you must deduce and add periods (full stops) where sentences naturally end. Always capitalize the first letter of the word immediately following a period.

Break the continuous text down into readable paragraphs of roughly 150 to 200 words. Additionally, if the topic clearly transitions, start a new paragraph.

Create a short, appropriate title/heading for every single paragraph based on what that section is about. Place the title right above its matching paragraph.

Find the most important keywords, phrases, or main ideas inside each formatted paragraph and make them bold to make it easier to read.

Return only the final formatted text. Do not add any extra introductory or concluding sentences before or after the text.
"""

# Define a mini-prompt just for cleaning up the Title
TITLE_PROMPT = """
If the following text contains any Hindi/Devanagari script, transliterate it completely into Hinglish (Hindi words written using the English alphabet). 
If it is already entirely in English, return it exactly as is. 
Only return the transformed title text. Do not add quotes, introductions, or extra punctuation.
"""

def add_page_number(run):
    """Injects Word XML field codes to auto-generate page numbers."""
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')

    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE"

    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')

    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')

    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)
    run._r.append(fldChar3)

def get_clipboard_text():
    """Silently grabs whatever text is currently copied to the clipboard."""
    root = tk.Tk()
    root.withdraw()
    try:
        text = root.clipboard_get()
    except tk.TclError:
        text = ""
    root.destroy()
    return text

def extract_video_id(value: str) -> str:
    value = value.strip()
    parsed = urlparse(value)
    if parsed.scheme in ("http", "https"):
        if parsed.netloc in ("youtu.be", "www.youtu.be"):
            return parsed.path.lstrip("/")
        qs = parse_qs(parsed.query)
        if "v" in qs:
            return qs["v"][0]
        m = re.search(r"/(shorts|embed)/([^/?&#]+)", parsed.path)
        if m:
            return m.group(2)
        return parsed.path.lstrip("/")
    return value

def clean_filename(name: str) -> str:
    return re.sub(r'[\\/*?:"<>|]', "", name).strip()

def safe_add_heading(doc, text, level):
    """Safely adds a heading, falling back to bold text if the style is missing."""
    try:
        doc.add_heading(text, level=level)
    except KeyError:
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = True

def setup_document():
    """Initializes and formats the main Word Document."""
    doc = Document()
    
    # 1. Set Custom Margins
    for section in doc.sections:
        section.top_margin = Inches(0)
        section.bottom_margin = Inches(0)
        section.left_margin = Inches(0.17)
        section.right_margin = Inches(4.0)
        
        # 2. Add Page Numbers to Footer Center
        footer = section.footer
        footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = footer_para.add_run()
        add_page_number(run)

    # 3. Shrink overall fonts
    style_normal = doc.styles['Normal']
    style_normal.font.size = Pt(8)
    
    try:
        doc.styles['Heading 1'].font.size = Pt(14)
        doc.styles['Heading 2'].font.size = Pt(12)
        doc.styles['Heading 3'].font.size = Pt(10)
    except KeyError:
        pass
        
    return doc

def append_to_doc(doc, formatted_text, title, channel, video_url):
    """Parses Gemini output and appends it to the master document."""
    safe_add_heading(doc, f"Source: {title}", level=1)
    
    meta_p = doc.add_paragraph()
    meta_p.add_run(f"Channel: {channel}\nLink: {video_url}").italic = True
    
    lines = formatted_text.split('\n')
    
    for line in lines:
        text = line.strip()
        if not text:
            continue 
            
        if text.startswith("### "):
            safe_add_heading(doc, text[4:].strip(), level=3)
        elif text.startswith("## "):
            safe_add_heading(doc, text[3:].strip(), level=2)
        elif text.startswith("# "):
            safe_add_heading(doc, text[2:].strip(), level=1)
        else:
            para = doc.add_paragraph()
            parts = text.split("**")
            
            for idx, part in enumerate(parts):
                if part: 
                    run = para.add_run(part)
                    if idx % 2 != 0:
                        run.bold = True
    
    # Add a page break after each video
    doc.add_page_break()

def get_playlist_info(playlist_url):
    """Uses yt-dlp to safely extract all video IDs from a playlist."""
    ydl_opts = {
        'extract_flat': True,
        'quiet': True,
        'ignoreerrors': True
    }
    with yt_dlp.YoutubeDL(ydl_opts) as ydl:
        info = ydl.extract_info(playlist_url, download=False)
        if not info:
            return "Unknown Playlist", []
            
        playlist_title = info.get('title', 'Unknown_Playlist')
        
        videos = []
        # 'entries' contains the list of videos
        if 'entries' in info:
            for entry in info['entries']:
                if entry and entry.get('id'):
                    videos.append({
                        'id': entry['id'],
                        'title': entry.get('title', 'Unknown Title'),
                        'channel': entry.get('uploader', 'Unknown Channel')
                    })
        return playlist_title, videos

def process_video(video_id, raw_title, channel):
    """Fetches and processes transcript for a single video."""
    print(f"\n▶ Processing: {raw_title}")
    
    try:
        title_response = model.generate_content([TITLE_PROMPT, raw_title])
        title = title_response.text.strip()
    except Exception:
        title = raw_title

    api = YouTubeTranscriptApi()
    
    try:
        transcript_list = api.list_transcripts(video_id)
        try:
            transcript_obj = transcript_list.find_transcript(['hi', 'en'])
        except:
            transcript_obj = transcript_list.find_generated_transcript(['hi', 'en'])
            
        raw_data = transcript_obj.fetch()
        text_lines = [entry['text'] for entry in raw_data]
        
    except Exception as transcript_err:
        try:
            transcript_obj = api.fetch(video_id, languages=['hi', 'en'])
            text_lines = [entry['text'] if isinstance(entry, dict) else entry.text for entry in transcript_obj]
        except Exception as e:
            print(f"  ❌ Failed to fetch transcript for {video_id}: {e}")
            return None, title

    raw_full_text = " ".join(text_lines)
    
    print("  🤖 Sending to Gemini...")
    try:
        response = model.generate_content([PROMPT, raw_full_text])
        return response.text, title
    except Exception as e:
        print(f"  ❌ Gemini API error: {e}")
        return None, title


def main():
    url_input = get_clipboard_text().strip()

    if "youtube.com" not in url_input and "youtu.be" not in url_input:
        print("❌ Error: No YouTube URL found in your clipboard.")
        input("\nPlease copy a YouTube link and try again. Press Enter to exit...")
        return

    print(f"📋 Read from clipboard: {url_input}\n")

    # Check if the URL is a playlist or a single video
    is_playlist = "list=" in url_input
    
    videos_to_process = []
    doc_title = "Single_Video"
    
    if is_playlist:
        print("📂 Playlist detected! Extracting videos... (this might take a few seconds)")
        playlist_title, videos = get_playlist_info(url_input)
        print(f"Found {len(videos)} videos in playlist: {playlist_title}")
        videos_to_process = videos
        doc_title = clean_filename(playlist_title)
    else:
        print("🎥 Single video detected.")
        video_id = extract_video_id(url_input)
        videos_to_process = [{'id': video_id, 'title': 'Single Video', 'channel': 'Unknown'}]

    if not videos_to_process:
        print("❌ No valid videos found.")
        return

    # Setup the Master Document
    doc = setup_document()
    
    base_dir = "Transcripts_YouTube"
    os.makedirs(base_dir, exist_ok=True)

    success_count = 0
    
    for idx, video in enumerate(videos_to_process):
        video_id = video['id']
        raw_title = video['title']
        channel = video['channel']
        video_url = f"https://www.youtube.com/watch?v={video_id}"
        
        # Process the video via Gemini
        formatted_text, final_title = process_video(video_id, raw_title, channel)
        
        if formatted_text:
            append_to_doc(doc, formatted_text, final_title, channel, video_url)
            success_count += 1
            
        # Add a short delay to avoid hitting Gemini's rate limits when looping fast
        if idx < len(videos_to_process) - 1:
            time.sleep(2)

    if success_count > 0:
        # Save the master document
        if not is_playlist:
            doc_title = clean_filename(final_title)
            
        doc_file_path = os.path.join(base_dir, f"{doc_title}_Master.docx")
        doc.save(doc_file_path)
        print(f"\n✅ SUCCESS! Saved {success_count} formatted transcripts to Master Word doc:")
        print(f"📁 {doc_file_path}")
    else:
        print("\n❌ Failed to process any transcripts.")

    input("\nPress Enter to close this window...")

if __name__ == "__main__":
    try:
        main()
    except Exception as e:
        print("\n" + "="*50)
        print("❌ A FATAL ERROR OCCURRED:")
        print("="*50)
        traceback.print_exc()
        print("="*50)
        input("\nPress Enter to close this window...")