import os
import argparse
from docx import Document

def format_markdown_in_docx(folder_path):
    folder_path = os.path.abspath(folder_path)

    if not os.path.isdir(folder_path):
        print(f"Error: The directory {folder_path} does not exist.")
        return

    # Find all docx files that look like our generated transcriptions
    docx_files = [f for f in os.listdir(folder_path) if f.endswith('_transcribed.docx')]
    
    if not docx_files:
        print(f"No transcribed .docx files found in {folder_path}.")
        return

    print(f"Found {len(docx_files)} document(s) to fix. Starting cleanup...\n")

    for filename in docx_files:
        file_path = os.path.join(folder_path, filename)
        print(f"Fixing: {filename}...")

        try:
            old_doc = Document(file_path)
            new_doc = Document()
            
            # Go through the old document
            for para in old_doc.paragraphs:
                # CRITICAL FIX: Split the giant paragraph by hidden line breaks
                lines = para.text.split('\n')
                
                for line in lines:
                    text = line.strip()
                    if not text:
                        continue # Skip empty lines
                        
                    # 1. Handle Headings
                    if text.startswith("### "):
                        new_doc.add_heading(text[4:].strip(), level=3)
                    elif text.startswith("## "):
                        new_doc.add_heading(text[3:].strip(), level=2)
                    elif text.startswith("# "):
                        new_doc.add_heading(text[2:].strip(), level=1)
                    
                    # 2. Handle Normal Paragraphs and Bold Text
                    else:
                        new_para = new_doc.add_paragraph()
                        parts = text.split("**")
                        
                        for i, part in enumerate(parts):
                            if part: 
                                run = new_para.add_run(part)
                                if i % 2 != 0:
                                    run.bold = True

            # Save and overwrite the messed up file
            new_doc.save(file_path)
            print("  Done.")

        except Exception as e:
            print(f"\nAn error occurred while formatting {filename}: {e}\n")

    print("\nAll files fixed successfully!")

if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Fix markdown formatting in transcribed Word documents.")
    parser.add_argument("folder_path", type=str, nargs='?', default=".", help="Path to the folder containing your .docx files.")
    
    args = parser.parse_args()
    format_markdown_in_docx(args.folder_path)