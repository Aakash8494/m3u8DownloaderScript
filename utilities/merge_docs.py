import os
import argparse
from docx import Document

def merge_transcriptions_continuous(folder_path, output_filename="All_Transcriptions_Combined.docx"):
    folder_path = os.path.abspath(folder_path)

    if not os.path.isdir(folder_path):
        print(f"Error: The directory {folder_path} does not exist.")
        return

    # Find only the transcribed files, sort them alphabetically
    docx_files = sorted([f for f in os.listdir(folder_path) if f.endswith('_transcribed.docx')])
    
    if not docx_files:
        print(f"No '_transcribed.docx' files found in {folder_path} to merge.")
        return

    print(f"Found {len(docx_files)} document(s) to merge continuously. Starting...\n")

    # Create the new master document
    master_doc = Document()
    master_doc.add_heading('Combined Audio Transcriptions', level=0)

    for i, filename in enumerate(docx_files):
        file_path = os.path.join(folder_path, filename)
        print(f"Adding: {filename}...")

        try:
            # Add a title for the audio file being appended
            clean_title = filename.replace('_transcribed.docx', '')
            master_doc.add_heading(f"Source: {clean_title}", level=1)
            
            sub_doc = Document(file_path)
            
            # Carefully copy paragraphs and their formatting (bold, styles)
            for para in sub_doc.paragraphs:
                new_para = master_doc.add_paragraph()
                new_para.style = para.style
                
                for run in para.runs:
                    new_run = new_para.add_run(run.text)
                    new_run.bold = run.bold
                    new_run.italic = run.italic
                    new_run.underline = run.underline

            # We added an empty line here instead of a page break for a little breathing room
            if i < len(docx_files) - 1:
                master_doc.add_paragraph()

        except Exception as e:
            print(f"  -> Error merging {filename}: {e}")

    # Save the final master document
    output_path = os.path.join(folder_path, output_filename)
    master_doc.save(output_path)
    print(f"\nSuccess! All files merged continuously into: {output_filename}")

if __name__ == "__main__":
    parser = argparse.ArgumentParser(description="Merge multiple Word documents into one continuous file.")
    parser.add_argument("folder_path", type=str, nargs='?', default=".", help="Path to the folder containing your .docx files.")
    parser.add_argument("--output", type=str, default="All_Transcriptions_Combined.docx", help="Name of the final merged file.")
    
    args = parser.parse_args()
    merge_transcriptions_continuous(args.folder_path, args.output)